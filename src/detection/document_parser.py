"""Document parsers for PDF, DOCX, TXT, and other formats."""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse various document formats into text for PII scanning."""

    SUPPORTED_EXTENSIONS = {".txt", ".log", ".csv", ".json", ".xml", ".pdf", ".docx", ".md", ".yaml", ".yml"}

    def parse(self, file_path: str | Path | None = None, content: bytes | None = None, filename: str = "") -> str:
        """Parse a document and return its text content."""
        if file_path:
            path = Path(file_path)
            ext = path.suffix.lower()
            filename = path.name
            with open(path, "rb") as f:
                content = f.read()
        elif content and filename:
            ext = Path(filename).suffix.lower()
        else:
            raise ValueError("Provide either file_path or both content and filename")

        if ext == ".pdf":
            return self._parse_pdf(content)
        elif ext == ".docx":
            return self._parse_docx(content)
        elif ext in (".txt", ".log", ".csv", ".md", ".yaml", ".yml", ".xml"):
            return self._parse_text(content)
        elif ext == ".json":
            return self._parse_text(content)
        else:
            logger.warning(f"Unsupported file type: {ext}, attempting plain text parse")
            return self._parse_text(content)

    def _parse_text(self, content: bytes) -> str:
        """Parse plain text content."""
        try:
            import chardet  # type: ignore[import-not-found]
            detected = chardet.detect(content)
            encoding = detected.get("encoding", "utf-8") or "utf-8"
            return content.decode(encoding)
        except Exception:
            return content.decode("utf-8", errors="replace")

    def _parse_pdf(self, content: bytes) -> str:
        """Parse PDF content."""
        try:
            from PyPDF2 import PdfReader  # type: ignore[import-not-found]
            reader = PdfReader(io.BytesIO(content))
            text_parts: list[str] = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.error("PyPDF2 not installed. Cannot parse PDF files.")
            return ""
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return ""

    def _parse_docx(self, content: bytes) -> str:
        """Parse DOCX content."""
        try:
            from docx import Document  # type: ignore[import-not-found]
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            return "\n".join(paragraphs)
        except ImportError:
            logger.error("python-docx not installed. Cannot parse DOCX files.")
            return ""
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return ""

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return ext in cls.SUPPORTED_EXTENSIONS
